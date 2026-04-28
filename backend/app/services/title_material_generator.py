from datetime import datetime
from io import BytesIO
from typing import Dict

from sqlalchemy.orm import Session

from backend.app.db.models import TeacherStats, User


class TitleMaterialGenerator:
    def __init__(self, db: Session, teacher_id: int) -> None:
        self.db = db
        self.teacher_id = teacher_id

    def generate_docx(self, include_appendix: bool = True) -> BytesIO:
        from docx import Document

        teacher = self._get_teacher()
        stats = self._get_stats()
        dims = stats.five_dimensions or {}

        doc = Document()
        title = doc.add_heading(f"{teacher.full_name or teacher.username} 职称申报材料", level=1)
        title.alignment = 1

        doc.add_heading("一、基本信息", level=2)
        doc.add_paragraph(f"姓名：{teacher.full_name or teacher.username}")
        doc.add_paragraph(f"学校：{teacher.school_name or '青山村小学'}")
        doc.add_paragraph("申报职称：待定")

        doc.add_heading("二、教学工作量", level=2)
        doc.add_paragraph(f"本学年备课次数：{stats.total_lesson_plans} 次")
        doc.add_paragraph(f"作业批改批次：{stats.total_homework_batches} 次")

        doc.add_heading("三、专业发展", level=2)
        doc.add_paragraph(f"微课分析与研修次数：{stats.total_micro_analysis} 次")
        doc.add_paragraph(f"教学能力雷达图评分：{self._dimension_summary(dims)}")

        doc.add_heading("四、家校沟通", level=2)
        doc.add_paragraph(f"家访记录次数：{stats.total_visit_records} 次")

        doc.add_heading("五、互助贡献", level=2)
        doc.add_paragraph(f"论坛发帖/回复次数：{stats.total_forum_posts} 次")

        if include_appendix:
            doc.add_heading("附录：核心指标摘要", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "指标"
            table.rows[0].cells[1].text = "数值"
            for key, value in [
                ("家校沟通", dims.get("communication", 0)),
                ("课堂引入", dims.get("intro", 0)),
                ("知识讲解", dims.get("explain", 0)),
                ("作业质量", dims.get("homework_quality", 0)),
                ("互动设计", dims.get("interaction", 0)),
            ]:
                row = table.add_row().cells
                row[0].text = key
                row[1].text = str(value)

        doc.add_paragraph(f"\n填表日期：{datetime.now().strftime('%Y年%m月%d日')}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_pdf(self, include_appendix: bool = True) -> BytesIO:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas

        teacher = self._get_teacher()
        stats = self._get_stats()
        dims = stats.five_dimensions or {}

        buffer = BytesIO()
        canv = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 50

        try:
            pdfmetrics.registerFont(TTFont("SimSun", "simsun.ttc"))
            canv.setFont("SimSun", 12)
        except Exception:
            canv.setFont("Helvetica", 12)

        lines = [
            f"{teacher.full_name or teacher.username} 职称申报材料",
            "",
            "一、基本信息",
            f"姓名：{teacher.full_name or teacher.username}",
            f"学校：{teacher.school_name or '青山村小学'}",
            "申报职称：待定",
            "",
            "二、教学工作量",
            f"本学年备课次数：{stats.total_lesson_plans} 次",
            f"作业批改批次：{stats.total_homework_batches} 次",
            "",
            "三、专业发展",
            f"微课分析与研修次数：{stats.total_micro_analysis} 次",
            f"教学能力雷达图评分：{self._dimension_summary(dims)}",
            "",
            "四、家校沟通",
            f"家访记录次数：{stats.total_visit_records} 次",
            "",
            "五、互助贡献",
            f"论坛发帖/回复次数：{stats.total_forum_posts} 次",
        ]
        if include_appendix:
            lines.extend(
                [
                    "",
                    "附录：核心指标摘要",
                    f"家校沟通：{dims.get('communication', 0)}",
                    f"课堂引入：{dims.get('intro', 0)}",
                    f"知识讲解：{dims.get('explain', 0)}",
                    f"作业质量：{dims.get('homework_quality', 0)}",
                    f"互动设计：{dims.get('interaction', 0)}",
                ]
            )
        lines.append("")
        lines.append(f"填表日期：{datetime.now().strftime('%Y年%m月%d日')}")

        for line in lines:
            canv.drawString(50, y, line)
            y -= 20
            if y < 40:
                canv.showPage()
                y = height - 50
                try:
                    canv.setFont("SimSun", 12)
                except Exception:
                    canv.setFont("Helvetica", 12)

        canv.save()
        buffer.seek(0)
        return buffer

    def _get_teacher(self) -> User:
        teacher = self.db.query(User).filter(User.id == self.teacher_id).first()
        if not teacher:
            raise ValueError("教师不存在")
        return teacher

    def _get_stats(self) -> TeacherStats:
        stats = (
            self.db.query(TeacherStats)
            .filter(TeacherStats.teacher_id == self.teacher_id)
            .first()
        )
        if stats:
            return stats
        return TeacherStats(
            teacher_id=self.teacher_id,
            total_lesson_plans=0,
            total_micro_analysis=0,
            total_visit_records=0,
            total_homework_batches=0,
            total_forum_posts=0,
            five_dimensions={},
        )

    @staticmethod
    def _dimension_summary(dims: Dict) -> str:
        return (
            f"家校沟通{dims.get('communication', 0)}分，"
            f"课堂引入{dims.get('intro', 0)}分，"
            f"知识讲解{dims.get('explain', 0)}分，"
            f"作业质量{dims.get('homework_quality', 0)}分，"
            f"互动设计{dims.get('interaction', 0)}分"
        )
